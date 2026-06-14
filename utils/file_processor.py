"""File processing utilities for extracting text from various file types"""
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Tuple, Optional


class FileProcessor:
    """Process different file types and extract text content"""
    
    @staticmethod
    def process_file(file_path: str) -> Tuple[str, int, dict]:
        """
        Process a file and extract text content
        
        Returns:
            Tuple of (extracted_text, word_count, metadata)
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            return FileProcessor._process_pdf(file_path)
        elif extension == ".docx":
            return FileProcessor._process_docx(file_path)
        elif extension in [".txt", ".md"]:
            return FileProcessor._process_text(file_path)
        elif extension in [".csv", ".xlsx"]:
            return FileProcessor._process_spreadsheet(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    @staticmethod
    def _process_pdf(file_path: Path) -> Tuple[str, int, dict]:
        """Extract text from PDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            metadata = {
                "page_count": len(doc),
                "author": doc.metadata.get("author", ""),
                "title": doc.metadata.get("title", ""),
                "subject": doc.metadata.get("subject", "")
            }
            
            for page in doc:
                text += page.get_text()
            
            doc.close()
            
            word_count = len(text.split())
            return text, word_count, metadata
        
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    @staticmethod
    def _process_docx(file_path: Path) -> Tuple[str, int, dict]:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            if tables_text:
                text += "\n\nTables:\n" + "\n".join(tables_text)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or ""
            }
            
            word_count = len(text.split())
            return text, word_count, metadata
        
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    @staticmethod
    def _process_text(file_path: Path) -> Tuple[str, int, dict]:
        """Extract text from TXT or MD files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            word_count = len(text.split())
            metadata = {
                "line_count": len(text.splitlines()),
                "encoding": "utf-8"
            }
            
            return text, word_count, metadata
        
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    @staticmethod
    def _process_spreadsheet(file_path: Path) -> Tuple[str, int, dict]:
        """Extract text from CSV or XLSX files"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == ".csv":
                df = pd.read_csv(file_path)
            else:  # .xlsx
                df = pd.read_excel(file_path, sheet_name=None)
                
                # Combine all sheets
                if isinstance(df, dict):
                    combined_text = []
                    for sheet_name, sheet_df in df.items():
                        combined_text.append(f"Sheet: {sheet_name}")
                        combined_text.append(sheet_df.to_string(index=False))
                    text = "\n\n".join(combined_text)
                    
                    metadata = {
                        "sheet_count": len(df),
                        "row_count": sum(len(sheet_df) for sheet_df in df.values()),
                        "column_count": sum(len(sheet_df.columns) for sheet_df in df.values())
                    }
                else:
                    text = df.to_string(index=False)
                    metadata = {
                        "row_count": len(df),
                        "column_count": len(df.columns)
                    }
                
                word_count = len(text.split())
                return text, word_count, metadata
            
            # For CSV
            text = df.to_string(index=False)
            metadata = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns)
            }
            
            word_count = len(text.split())
            return text, word_count, metadata
        
        except Exception as e:
            raise Exception(f"Error processing spreadsheet: {str(e)}")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """
    Split text into overlapping chunks
    
    Args:
        text: Input text
        chunk_size: Maximum chunk size in characters
        overlap: Overlap between chunks
    
    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        prev_start = start
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings
            sentence_endings = ['. ', '! ', '? ', '\n']
            best_break = end
            
            for ending in sentence_endings:
                pos = text.rfind(ending, start, end)
                if pos > start + overlap:
                    best_break = pos + len(ending)
                    break
            
            end = best_break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        if start <= prev_start:
            start = prev_start + chunk_size - overlap
            if start <= prev_start:
                start = prev_start + 1
    
    return chunks
